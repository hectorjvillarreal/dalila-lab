#!/usr/bin/env python3
"""
Revision script for "Is Health a Blessing?" — OLG Health Paper
Applies all four passes from OLGHealth_RevisionPlan.md

Output: OLGHealth_IDB_v1.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import re

doc = Document('OLGHealth_source.docx')

# ============================================================================
# PASS 1: ABSTRACT REWRITE
# ============================================================================

# The abstract is paragraphs 6 and 7 (paragraph 5 is the "Abstract" label)
# We replace both with a single new abstract paragraph

NEW_ABSTRACT = (
    "Middle-income economies across Latin America are aging at unprecedented speed "
    "— Mexico's total fertility rate reached 1.55 in 2024, below the United States "
    "at 1.62 — yet their fiscal systems remain unprepared for the health expenditure "
    "consequences. Unlike high-income countries with extensive social insurance, "
    "economies like Mexico combine low public health spending (626 USD PPP per capita, "
    "half the regional average), high out-of-pocket exposure, and a chronic disease "
    "burden led by diabetes affecting 16.9% of the population. Existing overlapping "
    "generations (OLG) models addressing health and aging are overwhelmingly calibrated "
    "to high-income contexts, leaving the middle-income, high-informality environment "
    "underexplored. We develop a dynamic stochastic partial equilibrium OLG model in "
    "which health status operates as a binary state variable affecting utility, survival "
    "probability, labor productivity, and out-of-pocket expenditure, and calibrate it "
    "to Mexico using ENOE, CONAPO, CIEP, and OECD data. Three findings stand out: "
    "consumption gaps of 40–70% between healthy and unhealthy individuals during working "
    "years; a 78% welfare loss for a 30-year-old worker who develops poor health; and "
    "income stagnation beginning at age 30 — precisely when health expenditure needs "
    "accelerate — departing sharply from standard life-cycle predictions. These results "
    "demonstrate that health is not a sectoral expenditure item but a core determinant "
    "of macroeconomic stability and intergenerational equity. In economies aging before "
    "reaching high-income status, insufficient health investment today depresses "
    "productivity, weakens fiscal capacity, and risks crowding out investment in younger "
    "cohorts — accelerating the fiscal stress of demographic transition."
)

# Clear paragraph 6 (first abstract paragraph) and set new text
p6 = doc.paragraphs[6]
# Preserve the style and formatting
for run in p6.runs:
    run.text = ""
if p6.runs:
    p6.runs[0].text = NEW_ABSTRACT
else:
    p6.add_run(NEW_ABSTRACT)

# Clear paragraph 7 (second abstract paragraph) — content now merged into p6
p7 = doc.paragraphs[7]
for run in p7.runs:
    run.text = ""

# Update keywords (paragraph 9)
NEW_KEYWORDS = (
    "Keywords: Macroeconomic equilibrium; Health; Aging; Out-of-pocket expenditures; "
    "Overlapping generations; Mexico; Demographic transition; Fiscal sustainability; "
    "Middle-income economies; Latin America"
)
p9 = doc.paragraphs[9]
for run in p9.runs:
    run.text = ""
if p9.runs:
    p9.runs[0].text = NEW_KEYWORDS
else:
    p9.add_run(NEW_KEYWORDS)

print("PASS 1 complete: Abstract rewritten")

# ============================================================================
# PASS 2: INTRODUCTION RESTRUCTURE
# ============================================================================

# Paragraph 18 (first intro paragraph) — replace with demographic moment opening
NEW_INTRO_P1 = (
    "The global fertility transition has entered a new phase. For the first time in "
    "human history, the world's total fertility rate fell below replacement in 2023, "
    "and the decline is accelerating most rapidly in middle-income economies. Mexico's "
    "total fertility rate reached 1.55 in 2024, falling below the United States at 1.62 "
    "— an unprecedented reversal for a country that has not yet attained high-income "
    "status. This places Mexico squarely in the \"aging before wealth\" trap: the "
    "demographic pressures traditionally associated with Japan, Germany, or South Korea "
    "are arriving decades earlier in the development trajectory, without the fiscal "
    "margins those economies built during their high-growth periods. The consequences "
    "for health expenditure are immediate. As populations age, health spending pressure "
    "intensifies precisely as the contribution base narrows and labor income growth "
    "slows. Yet most overlapping generations (OLG) models addressing the macroeconomic "
    "effects of health and aging are calibrated to high-income economies with extensive "
    "social insurance — not to the institutional and epidemiological context of Latin "
    "America."
)

p18 = doc.paragraphs[18]
for run in p18.runs:
    run.text = ""
if p18.runs:
    p18.runs[0].text = NEW_INTRO_P1
else:
    p18.add_run(NEW_INTRO_P1)

# Paragraph 19 — revise to foreground CIEP data and health-fiscal nexus
NEW_INTRO_P2 = (
    "The health-fiscal nexus in middle-income economies operates through channels "
    "that existing literature has not fully captured. In Mexico, per capita public "
    "health expenditure stands at 626 USD PPP — half of what Chile, Colombia, Costa "
    "Rica, or Panama allocate, and comparable to South Africa (OECD, 2023). This "
    "chronic underinvestment coexists with one of the highest chronic disease burdens "
    "in the region: 16.9% of Mexico's population lives with diabetes, compared to an "
    "OECD average of 7% (OECD, 2023). Diabetes prevalence at this scale is not merely "
    "an epidemiological fact — it is a fiscal time bomb interacting with demographic "
    "aging. The combination of low public spending and high out-of-pocket exposure means "
    "that health shocks transmit directly to household balance sheets, depressing "
    "consumption, eroding savings, and increasing dependence on government transfers "
    "(CIEP, 2018; Haakenstad et al., 2023). Analyzing these variables in isolation has "
    "systematically underestimated the macroeconomic impact of health conditions on "
    "national development."
)

p19 = doc.paragraphs[19]
for run in p19.runs:
    run.text = ""
if p19.runs:
    p19.runs[0].text = NEW_INTRO_P2
else:
    p19.add_run(NEW_INTRO_P2)

# Paragraph 20 — revise to pull gap statement forward
NEW_INTRO_P3 = (
    "A growing body of OLG literature addresses the macroeconomic consequences of "
    "health over the life cycle. De Nardi, French, and Jones (2010) establish medical "
    "expenditure risk as a dominant driver of late-life saving behavior. Fioroni (2010) "
    "and Dalgaard and Strulik (2014) endogenize health investment and longevity within "
    "life-cycle settings. More recent contributions incorporate health transitions into "
    "general equilibrium frameworks with demographic dynamics (Franković and Wrzaczek, "
    "2020; Lim, 2020). However, this literature is overwhelmingly calibrated to "
    "high-income economies — the United States, Japan, Germany, South Korea — where "
    "extensive public health insurance and well-developed social safety nets mediate "
    "the household-level impact of health shocks. The middle-income context, "
    "characterized by low public health spending, high informality, limited social "
    "insurance, and early income stagnation, remains largely unexplored in quantitative "
    "macroeconomic modeling."
)

p20 = doc.paragraphs[20]
for run in p20.runs:
    run.text = ""
if p20.runs:
    p20.runs[0].text = NEW_INTRO_P3
else:
    p20.add_run(NEW_INTRO_P3)

# Paragraph 21 — contribution statement
NEW_INTRO_P4 = (
    "This paper makes three contributions to the literature. First, methodologically, "
    "we develop a dynamic stochastic partial equilibrium OLG model in which health "
    "status operates as a binary state variable affecting utility, survival probability, "
    "labor productivity, and out-of-pocket expenditure simultaneously — calibrated to "
    "Mexico using ENOE labor survey data, CONAPO demographic projections, CIEP fiscal "
    "data, and OECD health statistics. Second, empirically, we provide quantitative "
    "evidence of consumption gaps, welfare losses, and early-life income stagnation "
    "patterns specific to a middle-income Latin American economy, demonstrating that "
    "the macroeconomic effects of poor health are not marginal but transformational. "
    "Third, from a policy perspective, we show that fiscal stress from demographic "
    "aging is not confined to pension systems: health conditions are an underestimated "
    "driver of long-run fiscal imbalance, operating through lower tax bases, higher "
    "transfer demands, and rising public and private health expenditures simultaneously."
)

p21 = doc.paragraphs[21]
for run in p21.runs:
    run.text = ""
if p21.runs:
    p21.runs[0].text = NEW_INTRO_P4
else:
    p21.add_run(NEW_INTRO_P4)

# Paragraph 23 — structure paragraph, light edit
NEW_INTRO_P5 = (
    "The remainder of the paper is organized as follows. Section 2 presents the "
    "partial equilibrium OLG model with health heterogeneity and describes the related "
    "literature. Section 3 details the calibration strategy and parametrization for "
    "Mexico. Section 4 presents the quantitative results. Section 5 discusses the "
    "findings and their implications for fiscal policy and the demographic transition. "
    "An appendix documents the evolving burden of disease in Costa Rica, Mexico, and "
    "Panama using Global Burden of Disease data, providing the epidemiological "
    "foundation for the health dynamics modeled in this paper."
)

p23 = doc.paragraphs[23]
for run in p23.runs:
    run.text = ""
if p23.runs:
    p23.runs[0].text = NEW_INTRO_P5
else:
    p23.add_run(NEW_INTRO_P5)

print("PASS 2 complete: Introduction restructured")

# ============================================================================
# PASS 3: DISCUSSION SECTION RESTRUCTURE
# ============================================================================

# Paragraphs 345-351 are currently Heading 2 with full paragraph text as the heading.
# We need to:
#   1. Change each to Normal style with the content as body text
#   2. Insert proper short headings before each

# The discussion heading mapping:
discussion_sections = [
    (345, "4.1 Health as a Macroeconomic State Variable",
     "The results of this analysis, based on a dynamic and stochastic "
     "overlapping-generations model in partial equilibrium, highlight the central role "
     "of population health in shaping macroeconomic outcomes over the life cycle. Poor "
     "health conditions reduce labor productivity, depress income trajectories, and "
     "generate persistent consumption gaps between healthy and unhealthy individuals. "
     "These effects materialize not only at older ages — as standard life-cycle models "
     "would predict — but significantly earlier, altering saving behavior, increasing "
     "household indebtedness, and raising concerns about long-term fiscal and "
     "intergenerational sustainability. The model demonstrates that health status "
     "functions as a macroeconomic state variable with consequences comparable in "
     "magnitude to skill heterogeneity, challenging the traditional treatment of health "
     "as a sectoral concern separate from core macroeconomic dynamics."),

    (346, "4.2 The Compounding Mechanism",
     "A key insight of the model is that health operates as a compounding mechanism "
     "across the life cycle. Through its simultaneous effects on utility, survival "
     "probabilities, labor productivity, and out-of-pocket expenditures, poor health "
     "weakens households' earning capacity while increasing their expenditure needs. "
     "This dual channel generates a reinforcing cycle: lower income reduces the capacity "
     "to invest in health and accumulate precautionary savings, while deteriorating "
     "health further depresses productivity and raises medical costs. As a result, "
     "individuals in poor health face substantially lower lifetime consumption and "
     "asset accumulation, even when government transfers provide a minimum consumption "
     "floor. The divergence in consumption paths between individuals in good and poor "
     "health is particularly pronounced among lower-skilled households, where the "
     "productivity penalty from poor health is larger and the capacity to self-insure "
     "through savings is more limited."),

    (347, "4.3 Early-Life Income Stagnation: The Mexico Case",
     "The Mexican calibration illustrates these dynamics with particular clarity. "
     "Unlike the reference model, in which income continues to rise until middle age, "
     "the simulations calibrated for Mexico show that income levels stagnate around "
     "age 30. This stagnation coincides precisely with the age at which health-related "
     "expenditure needs begin to increase relative to income — a timing that has "
     "profound consequences for household financial resilience. Consequently, households "
     "begin relying on borrowing or transfers much earlier in the life cycle to smooth "
     "consumption in the presence of health shocks. This early onset of financial "
     "stress contrasts sharply with standard life-cycle predictions, where health-related "
     "borrowing typically emerges closer to retirement age. The finding suggests that "
     "in middle-income economies, the interaction between health expenditure timing "
     "and income dynamics creates vulnerabilities that models calibrated to high-income "
     "contexts do not capture."),

    (348, "4.4 Institutional and Epidemiological Context",
     "These results acquire additional force when situated within Mexico's institutional "
     "and epidemiological context. The country combines persistently low public health "
     "expenditure — 626 USD PPP per capita, half the regional average — with one of "
     "the highest chronic disease burdens in Latin America. Diabetes alone affects "
     "16.9% of the population, disproportionately impacting working-age adults whose "
     "productivity losses cascade through household income and savings. The model "
     "suggests that insufficient health investment today does not merely translate into "
     "higher medical costs at older ages; it actively undermines productivity and income "
     "generation during prime working years. This weakens households' ability to "
     "self-insure through savings and increases reliance on out-of-pocket spending and "
     "public transfers — precisely the channels through which fiscal stress accumulates. "
     "The partial equilibrium framework adopted here isolates these household-level "
     "mechanisms with clarity, though a general equilibrium extension would capture "
     "additional feedback through wages and interest rates."),

    (349, "4.5 Intergenerational Trade-offs",
     "From an intergenerational perspective, the findings point to an important "
     "trade-off embedded in fiscal allocation decisions. As health-related needs among "
     "adults and the elderly grow, greater fiscal resources are required to prevent "
     "sharp declines in consumption and welfare among affected cohorts. However, if "
     "these resources are mobilized without expanding overall fiscal space — a binding "
     "constraint in most middle-income economies — they may crowd out investment in "
     "the health and human capital of younger cohorts. Such dynamics risk perpetuating "
     "a cycle in which poor health outcomes today reduce productivity and fiscal "
     "capacity tomorrow, constraining future investment and growth. The 78% welfare "
     "loss experienced by a 30-year-old worker who develops poor health quantifies "
     "the magnitude of this intergenerational transmission: today's health failures "
     "become tomorrow's fiscal burdens."),

    (350, "4.6 Health Investment as Fiscal Strategy",
     "The analysis underscores that health should not be viewed solely as a sectoral "
     "expenditure item but as a core determinant of macroeconomic stability and "
     "intergenerational equity. In contexts like Mexico, where demographic aging "
     "interacts with low public health spending and high out-of-pocket exposure, "
     "neglecting the macroeconomic role of health amplifies fiscal vulnerabilities "
     "and deepens inequality over the life cycle. The consumption gaps of 40–70% "
     "documented in this paper are not residual effects but first-order macroeconomic "
     "distortions with implications for aggregate demand, savings rates, and the "
     "distribution of welfare across cohorts. Strengthening health investment and "
     "reducing households' exposure to health-related financial risk emerge not as "
     "social policy aspirations but as central components of a sustainable long-term "
     "fiscal strategy — a conclusion reinforced by the early-life timing of income "
     "stagnation documented in the Mexico calibration."),

    (351, "4.7 Demographic Transition and Fiscal Imbalance",
     "Beyond household-level dynamics, the results speak directly to the interaction "
     "between fiscal sustainability and the demographic transition — the central "
     "analytical challenge for economies aging before reaching high-income status. "
     "Population aging alters the age composition of health expenditure toward older "
     "cohorts, while declining fertility — Mexico's TFR of 1.55 now sits below "
     "replacement and below the United States — limits the growth of the contribution "
     "base and future labor income. In this context, health-related shocks that depress "
     "productivity and income at earlier ages amplify fiscal pressures through multiple "
     "channels simultaneously: lower tax bases from reduced labor productivity, higher "
     "demand for transfers as households exhaust savings earlier, and rising health-"
     "related public and private expenditures as chronic disease prevalence increases "
     "with population aging. The model highlights that fiscal stress associated with "
     "aging is not confined to pension systems alone; health conditions act as an "
     "additional — and often underestimated — driver of long-run fiscal imbalance. "
     "For economies like Mexico, where the narrowing contribution base coincides with "
     "the income stagnation documented at age 30, the fiscal arithmetic of demographic "
     "transition is more adverse than pension-focused analyses suggest. Ignoring the "
     "macroeconomic consequences of population health understates the true fiscal cost "
     "of demographic transition and risks policy responses that arrive too late or "
     "target too narrow a set of fiscal channels."),
]

for idx, heading, body in discussion_sections:
    p = doc.paragraphs[idx]
    # Change the style from Heading 2 to Heading 3 for the short title
    # and write the body as the paragraph text
    # Strategy: clear the heading text, put short heading, then use the paragraph
    # for body text.
    #
    # Since we can't easily insert paragraphs between existing ones in python-docx,
    # we'll use the heading paragraph for the section heading and repurpose its text.
    # The body will go into the same paragraph for now — but that's not clean.
    #
    # Better approach: use the existing Heading 2 paragraph for the short heading,
    # and replace its text. Then we need to put body text somewhere.
    # The original text IS the heading — we replace it entirely.
    #
    # We'll change the style to Normal and prefix with the bold heading.

    # Clear all runs
    for run in p.runs:
        run.text = ""

    # Change style to Normal
    p.style = doc.styles['Normal']

    # Add heading as bold run, then body as normal run
    if p.runs:
        p.runs[0].text = heading + "\n"
        p.runs[0].bold = True
    else:
        run_h = p.add_run(heading + "\n")
        run_h.bold = True

    run_body = p.add_run(body)
    run_body.bold = False

print("PASS 3 complete: Discussion section restructured")

# ============================================================================
# PASS 4: MINOR EDITS
# ============================================================================

# 4a. Fix Dougherty reference (paragraph 360)
p360 = doc.paragraphs[360]
old_ref = p360.text
new_ref = old_ref.replace(
    "Dougherty, S., , P., & Lorenzoni, L.",
    "Dougherty, S., Lorenzoni, L., & Pisu, P."
)
for run in p360.runs:
    run.text = ""
if p360.runs:
    p360.runs[0].text = new_ref
else:
    p360.add_run(new_ref)

# 4b. Add appendix framing (paragraph 377 is "GBD: Costa Rica" — insert before it)
# We'll modify paragraph 376 area. Paragraph 376 is Heading 1 "Appendix".
# Paragraph 377 is "GBD: Costa Rica". We need to add framing text between them.
# Since inserting paragraphs is complex in python-docx, we'll prepend to paragraph 377.

p377 = doc.paragraphs[377]
APPENDIX_FRAMING = (
    "The following appendix documents the evolving burden of disease in three Latin "
    "American countries using Global Burden of Disease (GBD) data. Costa Rica, Mexico, "
    "and Panama bracket the region's epidemiological transition: Costa Rica represents "
    "an advanced transition with relatively high public health spending; Mexico occupies "
    "a middle position with low public spending and exceptionally high diabetes "
    "prevalence; and Panama represents an intermediate case. Together, these three "
    "profiles illustrate the structural persistence of chronic noncommunicable diseases "
    "across the region and strengthen the generalizability of the health dynamics "
    "modeled in this paper.\n\n"
)

# Prepend the framing text
original_377 = p377.text
for run in p377.runs:
    run.text = ""
if p377.runs:
    p377.runs[0].text = APPENDIX_FRAMING + original_377
else:
    p377.add_run(APPENDIX_FRAMING + original_377)

# 4c. The partial equilibrium defense is already present in the Methods and was
# added in the discussion section 4.4 revision above. No additional edit needed.

print("PASS 4 complete: Minor edits applied")

# ============================================================================
# SAVE
# ============================================================================

output_path = 'OLGHealth_IDB_v1.docx'
doc.save(output_path)
print(f"\nSaved: {output_path}")
print("All four passes complete.")
